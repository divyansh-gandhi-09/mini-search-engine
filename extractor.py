from fastapi import FastAPI, UploadFile, HTTPException, File
from fastapi.middleware.cors import CORSMiddleware
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
import traceback
import logging
import re
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
import asyncio
from typing import List, Dict, Any
import sys

# Configure tesseract path
pytesseract.pytesseract.tesseract_cmd = r"C:\Users\HP\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"

# Optional imports with detailed error messages
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError as e:
    DOCX_AVAILABLE = False
    print(f"⚠️ DOCX support unavailable: {e}")

try:
    import openpyxl
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError as e:
    EXCEL_AVAILABLE = False
    print(f"⚠️ Excel support unavailable: {e}")

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError as e:
    HTML_AVAILABLE = False
    print(f"⚠️ HTML parsing unavailable: {e}")

#  CHECK CRITICAL DEPENDENCIES
try:
    import fitz
    print(f" PyMuPDF (fitz) version: {fitz.version}")
except ImportError as e:
    print(f" CRITICAL: PyMuPDF not available: {e}")
    print("   Install with: pip install PyMuPDF")
    sys.exit(1)

try:
    from PIL import Image
    print(f" Pillow (PIL) available")
except ImportError as e:
    print(f" CRITICAL: Pillow not available: {e}")
    print("   Install with: pip install Pillow")
    sys.exit(1)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="Ultra-Fast Document Extractor",
    description="Parallel batch text extraction with multi-core processing",
    version="3.0.1"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Use all CPU cores
import multiprocessing as mp
MAX_WORKERS = max(1, mp.cpu_count() // 2)
logger.info(f" Initialized with {MAX_WORKERS} parallel workers")

SUPPORTED_EXTENSIONS = {
    'pdf': 'PDF documents',
    'png': 'PNG images', 'jpg': 'JPEG images', 'jpeg': 'JPEG images',
    'tiff': 'TIFF images', 'bmp': 'BMP images', 'gif': 'GIF images',
    'txt': 'Plain text', 'md': 'Markdown', 'csv': 'CSV',
    'json': 'JSON', 'log': 'Log files', 'xml': 'XML',
    'html': 'HTML', 'htm': 'HTML',
    'docx': 'Word documents',
    'xlsx': 'Excel files', 'xls': 'Excel files'
}

def clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace('\r\n', '\n').replace('\r', '\n')   # normalize line endings first
    text = re.sub(r'\n{3,}', '\n\n', text)                  # collapse triple+ newlines
    text = re.sub(r'[ \t]+', ' ', text)                     # collapse horizontal space only
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)  # strip control chars
    return text.strip()

def extract_from_pdf(content: bytes, filename: str) -> dict:
    """Extract text from PDF - optimized for speed"""
    try:
        logger.info(f"📄 Extracting PDF: {filename} ({len(content)} bytes)")
        
        #  Verify content is not empty
        if not content or len(content) < 100:
            logger.error(f" PDF too small or empty: {filename}")
            return {
                'success': False,
                'filename': filename,
                'error': 'PDF file is empty or corrupted'
            }
        
        #  Check if it's actually a PDF
        if b'%PDF' not in content[:1024]:  # Check first 1KB
            logger.error(f" Not a valid PDF: {filename}")
            return {'success': False, 'error': 'File is not a valid PDF'}
        
        pdf = fitz.open(stream=content, filetype="pdf")
        page_count = len(pdf)
        logger.info(f"  📖 PDF has {page_count} pages")
        
        text_parts = []
        
        for page_num, page in enumerate(pdf, 1):
            try:
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(page_text)
                    logger.debug(f"   Page {page_num}: {len(page_text)} chars")
                else:
                    logger.warning(f"  ⚠️ Page {page_num}: Empty")
            except Exception as e:
                logger.error(f"   Page {page_num} failed: {e}")
        
        pdf.close()
        
        full_text = clean_text("\n\n".join(text_parts))
        
        if not full_text:
            logger.info(f"No text layer found in {filename}, attempting OCR on {page_count} pages...")
            ocr_parts = []
            pdf2 = fitz.open(stream=content, filetype="pdf")
            for page in pdf2:
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                img = preprocess_image(img)   # see FIX-09
                page_text = pytesseract.image_to_string(img, config='--oem 3 --psm 3')
                if page_text.strip():
                    ocr_parts.append(page_text)
            pdf2.close()
            full_text = clean_text("\n\n".join(ocr_parts))
            if full_text:
                logger.info(f"OCR fallback succeeded: {len(full_text)} chars")
            if not full_text:
        # OCR also found nothing — truly empty or unreadable
                logger.error(f"❌ OCR fallback also failed for: {filename}")
                return {
                    'success': False,
                    'filename': filename,
                    'error': f'No text extractable from {page_count} pages (unreadable or image-only PDF)'
                 }
    
            logger.info(f" OCR fallback succeeded: {len(full_text)} chars from {filename}")
        
        logger.info(f" PDF extracted: {len(full_text)} chars from {len(text_parts)} pages")
        
        return {
            'text': full_text,
            'pages': page_count,
            'extraction_method': 'PyMuPDF',
            'success': True,
            'filename': filename
        }
        
    except Exception as e:
        error_msg = f"PDF extraction failed: {str(e)}"
        logger.error(f" {filename}: {error_msg}")
        logger.error(f"   Traceback: {traceback.format_exc()}")
        return {
            'success': False,
            'filename': filename,
            'error': error_msg,
            'traceback': traceback.format_exc()
        }
from PIL import ImageFilter, ImageEnhance, ImageOps

def preprocess_image(image: Image.Image) -> Image.Image:
    """Improve image quality before OCR"""
    image = image.convert('L')  # grayscale
    # Upscale if too small — Tesseract needs ~300 DPI
    if image.width < 1000:
        scale = 1000 / image.width
        image = image.resize(
            (int(image.width * scale), int(image.height * scale)),
            Image.LANCZOS
        )
    image = ImageEnhance.Contrast(image).enhance(2.0)
    image = image.filter(ImageFilter.SHARPEN)
    return image

def extract_from_image(content: bytes, filename: str) -> dict:
    """Extract text from images using OCR"""
    try:
        logger.info(f" Extracting image: {filename} ({len(content)} bytes)")
        
        if not content:
            return {
                'success': False,
                'filename': filename,
                'error': 'Empty image file'
            }
        
        image = Image.open(io.BytesIO(content))
        image = preprocess_image(image)  
        logger.info(f"  Image size: {image.size}, mode: {image.mode}")
        
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Fast OCR with minimal overhead
        text = pytesseract.image_to_string(image, config='--oem 3 --psm 3')
        
        cleaned_text = clean_text(text)
        logger.info(f" OCR extracted: {len(cleaned_text)} chars")
        
        return {
            'text': cleaned_text,
            'extraction_method': 'Tesseract',
            'success': True,
            'filename': filename
        }
    except Exception as e:
        error_msg = f"OCR failed: {str(e)}"
        logger.error(f" {filename}: {error_msg}")
        return {
            'success': False,
            'filename': filename,
            'error': error_msg
        }

def extract_from_docx(content: bytes, filename: str) -> dict:
    """Extract from Word documents"""
    if not DOCX_AVAILABLE:
        return {'success': False, 'filename': filename, 'error': 'DOCX support not available'}
    
    try:
        logger.info(f"📝 Extracting DOCX: {filename}")
        doc = docx.Document(io.BytesIO(content))
        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        cleaned_text = clean_text("\n".join(text_parts))
        logger.info(f" DOCX extracted: {len(cleaned_text)} chars")
        
        return {
            'text': cleaned_text,
            'extraction_method': 'python-docx',
            'success': True,
            'filename': filename
        }
    except Exception as e:
        logger.error(f" DOCX extraction failed for {filename}: {str(e)}")
        return {
            'success': False,
            'filename': filename,
            'error': str(e)
        }

def extract_from_excel(content: bytes, filename: str) -> dict:
    """Extract from Excel files"""
    if not EXCEL_AVAILABLE:
        return {'success': False, 'filename': filename, 'error': 'Excel support not available'}
    
    try:
        logger.info(f"📊 Extracting Excel: {filename}")
        excel_file = pd.ExcelFile(io.BytesIO(content))
        text_parts = []
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            if not df.empty:
                text_parts.append(df.to_string())
        
        cleaned_text = clean_text("\n\n".join(text_parts))
        logger.info(f" Excel extracted: {len(cleaned_text)} chars")
        
        return {
            'text': cleaned_text,
            'extraction_method': 'pandas',
            'success': True,
            'filename': filename
        }
    except Exception as e:
        logger.error(f"❌ Excel extraction failed for {filename}: {str(e)}")
        return {
            'success': False,
            'filename': filename,
            'error': str(e)
        }

def extract_from_html(content: bytes, filename: str) -> dict:
    """Extract from HTML"""
    try:
        logger.info(f"🌐 Extracting HTML: {filename}")
        
        if not HTML_AVAILABLE:
            text = content.decode('utf-8', errors='replace')
            text = re.sub(r'<[^>]+>', '', text)
            return {
                'text': clean_text(text),
                'extraction_method': 'regex',
                'success': True,
                'filename': filename
            }
        
        soup = BeautifulSoup(content, 'html.parser')
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text()
        cleaned_text = clean_text(text)
        logger.info(f" HTML extracted: {len(cleaned_text)} chars")
        
        return {
            'text': cleaned_text,
            'extraction_method': 'BeautifulSoup',
            'success': True,
            'filename': filename
        }
    except Exception as e:
        logger.error(f" HTML extraction failed for {filename}: {str(e)}")
        return {
            'success': False,
            'filename': filename,
            'error': str(e)
        }

def extract_from_text(content: bytes, filename: str) -> dict:
    """Extract from text files"""
    logger.info(f"📃 Extracting text: {filename}")
    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
    
    for encoding in encodings:
        try:
            text = content.decode(encoding)
            cleaned_text = clean_text(text)
            logger.info(f" Text extracted ({encoding}): {len(cleaned_text)} chars")
            return {
                'text': cleaned_text,
                'extraction_method': f'text-{encoding}',
                'success': True,
                'filename': filename
            }
        except UnicodeDecodeError:
            continue
    
    # Fallback
    text = content.decode('utf-8', errors='replace')
    cleaned_text = clean_text(text)
    logger.warning(f"⚠️ Text extracted with fallback: {len(cleaned_text)} chars")
    return {
        'text': cleaned_text,
        'extraction_method': 'text-fallback',
        'success': True,
        'filename': filename
    }

def extract_single_file(content: bytes, filename: str) -> dict:
    """Route to appropriate extractor - OPTIMIZED"""
    file_extension = Path(filename).suffix.lower().lstrip('.')
    
    logger.info(f"🔧 Processing: {filename} (extension: .{file_extension})")
    
    try:
        if file_extension == 'pdf':
            result = extract_from_pdf(content, filename)
        elif file_extension in ['png', 'jpg', 'jpeg', 'tiff', 'bmp', 'gif']:
            result = extract_from_image(content, filename)
        elif file_extension == 'docx':
            result = extract_from_docx(content, filename)
        elif file_extension in ['xlsx', 'xls']:
            result = extract_from_excel(content, filename)
        elif file_extension in ['html', 'htm']:
            result = extract_from_html(content, filename)
        elif file_extension in ['txt', 'md', 'csv', 'json', 'log', 'xml']:
            result = extract_from_text(content, filename)
        else:
            logger.warning(f"⚠️ Unknown extension .{file_extension}, trying text extraction")
            result = extract_from_text(content, filename)
        
        return result
        
    except Exception as e:
        error_msg = f"Extraction failed: {str(e)}"
        logger.error(f" {filename}: {error_msg}")
        logger.error(f"   Traceback: {traceback.format_exc()}")
        return {
            'success': False,
            'filename': filename,
            'error': error_msg,
            'traceback': traceback.format_exc()
        }


@app.post("/extract")
async def extract_text(file: UploadFile = File(...)):
    """Single file extraction endpoint"""
    logger.info(f" Received single file: {file.filename}")
    
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
    
    content = await file.read()
    logger.info(f"  Read {len(content)} bytes")
    
    if not content:
        raise HTTPException(status_code=400, detail="Empty file")
    
    result = extract_single_file(content, file.filename)
    
    if not result['success']:
        logger.error(f" Extraction failed: {result.get('error', 'Unknown error')}")
        raise HTTPException(status_code=500, detail=result['error'])
    
    return {
        "text": result['text'],
        "filename": file.filename,
        "extracted_chars": len(result['text']),
        "extraction_method": result.get('extraction_method', 'Unknown'),
        "success": True
    }

@app.post("/extract/batch")
async def extract_batch_parallel(files: List[UploadFile] = File(...)):
    """
    ⚡ Process multiple files in PARALLEL using all CPU cores
    """
    
    if len(files) > 10000:
        raise HTTPException(status_code=400, detail="Max 10000 files per batch")
    
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")
    
    logger.info(f"🚀 Batch extraction: {len(files)} files on {MAX_WORKERS} workers")
    
    # Read all files asynchronously (fast I/O)
    async def read_file(file: UploadFile):
        content = await file.read()
        logger.info(f"  📥 Read {file.filename}: {len(content)} bytes")
        return (content, file.filename)
    
    file_data = await asyncio.gather(*[read_file(f) for f in files])
    
    logger.info(f"📥 Read {len(file_data)} files, starting parallel extraction...")
    
    #  Extract ALL files in parallel with ThreadPoolExecutor
    results = []
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = {
            executor.submit(extract_single_file, content, filename): filename
            for content, filename in file_data
        }
        
        # Collect results as they complete
        completed = 0
        for future in as_completed(futures, timeout=180):
            filename = futures[future]
            try:
                result = future.result()
                results.append(result)
                completed += 1
                
                if result['success']:
                    logger.info(f"[{completed}/{len(files)}] {filename}: {len(result['text'])} chars")
                else:
                    logger.error(f" [{completed}/{len(files)}] {filename}: {result.get('error', 'Unknown')}")
                
                if completed % 50 == 0:
                    logger.info(f"📊 Progress: {completed}/{len(files)} completed")
                    
            except Exception as e:
                logger.error(f" Fatal error processing {filename}: {e}")
                logger.error(f"   Traceback: {traceback.format_exc()}")
                results.append({
                    'success': False,
                    'filename': filename,
                    'error': str(e),
                    'traceback': traceback.format_exc()
                })
    
    successful = [r for r in results if r.get('success')]
    failed = [r for r in results if not r.get('success')]
    
    logger.info(f"🎉 Batch complete: {len(successful)} | {len(failed)} ")
    
    # Log failed files for debugging
    if failed:
        logger.error(" Failed files:")
        for fail in failed:
            logger.error(f"   - {fail['filename']}: {fail.get('error', 'Unknown')}")
    
    return {
        'results': results,
        'total': len(files),
        'successful': len(successful),
        'failed': len(failed),
        'workers_used': MAX_WORKERS,
        'success': len(successful) > 0
    }

@app.get("/")
async def root():
    return {
        "service": "Document Text Extractor",
        "version": "3.0.1",
        "features": [
            "Parallel batch processing",
            "10+ file format support",
            "Detailed error logging",
            "Metadata extraction"
        ],
        "supported_formats": SUPPORTED_EXTENSIONS,
        "dependencies": {
            "pymupdf": "available" if 'fitz' in sys.modules else "missing",
            "pillow": "available" if 'PIL' in sys.modules else "missing",
            "docx": "available" if DOCX_AVAILABLE else "unavailable",
            "excel": "available" if EXCEL_AVAILABLE else "unavailable",
            "html": "available" if HTML_AVAILABLE else "unavailable"
        }
    }

@app.get("/health")
async def health():
    return {
        "status": "ok",
        "workers": MAX_WORKERS,
        "dependencies": {
            "pymupdf": "available" if 'fitz' in sys.modules else "missing",
            "pillow": "available" if 'PIL' in sys.modules else "missing",
            "docx": DOCX_AVAILABLE,
            "excel": EXCEL_AVAILABLE,
            "html": HTML_AVAILABLE
        }
    }

if __name__ == "__main__":
    import uvicorn
    print("=" * 60)
    print("Document Text Extractor v3.0.1 - Enhanced Error Logging")
    print("=" * 60)
    print(f" PyMuPDF: {fitz.version}")
    print(f" Workers: {MAX_WORKERS}")
    print(f"Supported formats: {len(SUPPORTED_EXTENSIONS)}")
    print("Starting on http://0.0.0.0:5000")
    print("=" * 60)
    uvicorn.run(app, host="0.0.0.0", port=5000, log_level="info")